import docx
import re
import os
import argparse
argparse = argparse.ArgumentParser('处理docx，并发现重复句子')

argparse.add_argument('-init', action='store_true', help=u'将当前目录内的docx以标点初始化为多行txt')
argparse.add_argument('--show_common', '-run', action='store_true',help=u'显示当前docx重复句')
path1 = "D:\\My Space\\diff1.docx"
output1 = 'output1.txt'
now_path = './'
args = argparse.parse_args()

def getall_docx():
    """
    输出当前目录内全部docx，并按标点转换为多行
    :return:
    """
    files = os.listdir('.')
    for file in files:
        if 'docx' in file:
            print(file)
            output = input("请输入输出文件名:")
            diffdoc(file,output)
def diffdoc(path1,output):
    """
    将docx内的句子分割成多行
    :param path1:
    :return:
    """
    #获取文档对象
    file1=docx.Document(path1)
    pattern = r',|\.|/|;|\'|`|\[|\]|<|>|\?|:|"|\{|\}|\~|!|@|#|\$|%|\^|&|\(|\)|-|=|\_|\+|，|。|、|；|‘|’|【|】|·|！| |…|（|）'
    #输出每一段的内容
    for para in file1.paragraphs:
        print(para.text)

    #输出段落编号及段落内容
    for i in range(len(file1.paragraphs)):
        file1_line = file1.paragraphs[i].text
        result_list = re.split(pattern, file1_line)
        for j in result_list:
            with open("./"+output, 'a') as f:
                f.write("{}\n".format(j))
def diff_file():
    files = os.listdir('.')
    for file in files:
        print(file)
    file1 = input('[提示]请输入要比对的文件1:')
    file2 = input('[提示]请输入要比对的文件2:')
    str1 = []
    file_1 = open(file1, "r")
    for line in file_1.readlines():
        str1.append(line.replace("\n", ""))

    str2 = []
    file_2 = open(file2, "r")
    for line in file_2.readlines():
        str2.append(line.replace("\n", ""))


    common = []
    for string in str1:
        if string in str2:
            common.append(string)
    print('重复字段有:\n{}')
    for i in common:
        print(i)


# diffdoc(path1)
# getall_docx()


if __name__ == '__main__':
    if args.init:
        getall_docx()
    if args.show_common:
        diff_file()